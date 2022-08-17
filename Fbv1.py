from bs4 import BeautifulSoup
import api
from tkinter import *
from tkinter import filedialog
from docx import Document
import docx
from tkinter import messagebox
import pyautogui
import re
import random

root = Tk()
root.geometry("600x850")

html1 = """If you have created a bootable USB drive, you might be wondering whether it will successfully initialize and boot. You don’t always have to restart your PC/laptop to find out. The following techniques can easily determine whether a USB drive is bootable or not in Windows 10/11. These include methods native to your Windows system as well as external recommended software.
<div>
<div>
<ul>
 	<li></li>
 	<li></li>
 	<li></li>
 	<li></li>
 	<li></li>
 	<li></li>
 	<li></li>
</ul>
</div>
</div>"""

class editContent():
    def __init__(self,html):
        self.html = html
    def tailContent(self):
        tail_sheet = '1. Tail Content!C2:E'
        funct_sheet = '1. Tail Content!P2:P15'
        tag_sheet = '1. Tail Content!Q2:Q15'
        tail_id = '1Z_Jy-8gX5lC0i_BoF4DnTpb7E2Ah9sypvV8eHDm-9LI'

        dataset = []
        funct = []
        tag = []
        def createTuple(x,y,z):
            yield (x,y,z)
        def createSgTuple(x):
            yield (x)
        for x in api.get_data(tail_id,tail_sheet):
            try:
                tuple_create = createTuple(x[0],x[1],x[2])
                for x in tuple_create:
                    dataset.append(x)
            except:
                pass
        for y in api.get_data(tail_id,funct_sheet):
            try:
                tuple_funct = createSgTuple(y[0])
            except:
                pass
            for x in tuple_funct:
                funct.append(x)
        for z in api.get_data(tail_id,tag_sheet):
            try:
                tag_Craete = createSgTuple(z[0])
            except:
                pass
            for x in tag_Craete:
                tag.append(x)
        tag = tuple(tag)
        funct = tuple(funct)
        dataset = tuple(dataset)
        def replaceTail(data):
            for x in tag:
                for y in funct:
                    for z in ['',' ',' ']:
                        self.html = str(self.html).replace("</"+x+">" + str(data).strip(),"</"+x+">")
                        self.html = str(self.html).replace("</"+x+">" + y + z + str(data).strip(),"</"+x+">" + y)
                        self.html = str(self.html).replace("</"+x+">" + str(data).strip() + y,"</"+x+">" + y)

            return self.html
        def replaceHead():
            for x in tag:
                for y in list(map(chr, range(97, 123))):
                    self.html = str(self.html).replace(y + "<" + x + ">",y + " <" + x + ">")
                    self.html = str(self.html).replace("</" + x + ">" + y," </" + x + "> " + y)

            return self.html
        def setPriority(data):
            for z in data:
                if z[1] == "Xoá bỏ":
                    replaceTail(z[0])
                else:
                    pass
            replaceHead()
            return self.html
        return setPriority(dataset)
    def removeTag(self):
        spreadsheet_id = '1Z_Jy-8gX5lC0i_BoF4DnTpb7E2Ah9sypvV8eHDm-9LI'
        sheet_tags = '2. Tag!C4:G'

        def cr_tuple(x,y,z,t,h):
            yield (x,y,z,t,h)
        dataset = []
        for x in api.get_data(spreadsheet_id,sheet_tags):
            try:
                tuple = cr_tuple(x[0],x[1],x[2],x[3],x[4])
                for x in tuple:
                    dataset.append(x)
            except:
                pass
        def setPriority():
            def change_tag(html):
                global funt
                html_content = []
                funt = []
                for data in dataset:
                    if data[4] == "Ưu tiên 5":
                        funt.append(data[2])
                    else:
                        pass

                for l in list(dict.fromkeys(html.splitlines())):
                    if any(x in l for x in funt):
                        html = html.replace(l,"")
                    else:
                        pass

                return html

            change_html = change_tag(self.html)
            soup = BeautifulSoup(change_html, 'html.parser')

            def pri1(data):
                if data[4] == "Ưu tiên 1":
                    for x in soup.select(data[0]):
                        x.extract()
                else:
                    pass
            def pri2(data):
                pass
            def pri3(data):
                if data[4] == "Ưu tiên 3":
                    for x in soup.findAll(data[0]):
                        try:
                            if data[1] != '' and data[2] in x[data[1]]:
                                x.extract()
                            elif data[1] == '' and data[2] in x.text:
                                x.extract()
                            else:
                                pass
                        except:
                            pass
            def pri4(data):
                if data[4] == "Ưu tiên 4" and data[0] == 'a':
                    for x in soup.findAll(data[0], href=True):
                        if x.text:
                            if data[1] == '' and data[2] in x.text:
                                x.unwrap()
                            elif data[1] != '' and data[2] in x[data[1]]:
                                x.unwrap()
                            elif len(x.text) > 80:
                                x.unwrap()
                            else:
                                pass
                    for x in soup.findAll(data[0], href=False):
                        if x.text:
                            x.unwrap()
                if data[4] == "Ưu tiên 4" and data[0] != 'a':
                    for x in soup.findAll(data[0]):
                        if x.text:
                            if data[1] == '' and data[2] in x.text:
                                x.unwrap()
                            elif data[1] != '' and data[2] in x[data[1]]:
                                x.unwrap()
                            else:
                                pass
                return soup
            def pri5(data):
                if data[4] == "Ưu tiên 5":
                    for x in soup.find_all(text=re.compile(data[2])):
                        try:
                            x.parent.next_sibling.replaceWith("")
                            x.parent.decompose()
                        except:
                            pass

            def remove_junk():
                delete_list = ['code', 'li', 'h2', 'h3', 'table', 'strong', 'picture', 'figure', 'img']
                keep_list = ['source', 'img','td','li']

                for x in soup.findAll():
                    if x.get_text(strip=True).strip() == "" and x.name not in keep_list:
                        if any(x.find(y) != None for y in delete_list):
                            pass
                        else:
                            x.decompose()
                    else:
                        pass

                long_tag = ['strong','em','code','b']
                except_long = ['sudo','""','C:','HKEY','$Windows','$','DISM','/','~','*','%','localappdata']
                for y in long_tag:
                    for x in soup.findAll(y):
                        if len(x.text) > 30 and not any(y in x.text for y in except_long):
                            x.unwrap()
                        else:
                            pass
                short_tag = ['table']
                for z in short_tag:
                    for x in soup.findAll(z):
                        if len(x.text) > 1500:
                            x.unwrap()
                        else:
                            pass


            def priority(data):
                pri1(data)
                pri2(data)
                pri3(data)
                pri4(data)

            for x in dataset:
                priority(x)
            remove_junk()

            funt = []
            for data in dataset:
                if data[4] == "Ưu tiên 5":
                    funt.append(data[2])
                else:
                    pass

                if data[4] == "Ưu tiên 2":
                    soup = str(soup).replace(data[2], "")
                elif data[4] == "Ưu tiên 6":
                    for l in list(dict.fromkeys(str(soup).splitlines())):
                        for x in l.split("."):
                            for y in x.split("."):
                                try:
                                    if data[0] in y and any(data[0].split()[0] in z for z in y.split(" ")):
                                        x = y.replace(data[0], data[2])
                                        soup = soup.replace(y, x)
                                    else:
                                        pass
                                except:
                                    pass
                else:
                    pass

            soup = str(soup).replace("<code>","<strong>")
            soup = str(soup).replace("</code>","</strong>")

            return soup
        return setPriority()
    def convertData(self):
        data_convert = ['strong','em']
        for x in data_convert:
            self.html = self.html.replace("<" + x + ">","<code>")
            self.html = self.html.replace("</" + x + ">","</code>")

        return self.html
    def removeJunk(self):
        spreadsheet_id = '1Z_Jy-8gX5lC0i_BoF4DnTpb7E2Ah9sypvV8eHDm-9LI'
        sheet_tags = '3. Delete!C4:C'

        dataset = []
        for x in api.get_data(spreadsheet_id,sheet_tags):
            try:
                dataset.append(x[0])
            except:
                pass
        for x in dataset:
            self.html = self.html.replace(x,"")

        return self.html

class GUI():
    def __init__(self):
        root.title("EDIT CONTENT")
    def main_execution(self):
        output = editContent().main()
    def main_GUI(self,mode = None):
        def chooseFile():
            root.filename = filedialog.askopenfilename(initialdir="/Volumes/Root/Webify",title="Choose the path",filetypes=(("Word file","*.doc"),("Word file","*.docx"),("Note file","*.txt"),("all files","*.*")))
            doc = docx.Document(root.filename)
            all_paras = doc.paragraphs
            content = []
            for x in all_paras:
                content.append(x.text)
            content_trans = "\n".join(content)
            a = editContent(content_trans).tailContent()
            b = editContent(a).removeTag()
            list_b = b.split("\n")
            document = Document()
            for x in list_b:
                document.add_paragraph(x)
            document.save(root.filename)

            messagebox.showinfo("Status message","Successfully transfer")

        def insertData():
            def insertValue(mode=None):
                global Label_test
                global c
                global current
                current = data_output.get("1.0", "end-1c")
                if mode == 'tailcontent':
                    a = editContent(str(current)).tailContent()
                    b = editContent(str(a)).removeTag()
                    c = editContent(str(b)).removeJunk()
                    data_output_out.insert(END,str(c))

                    response = messagebox.showinfo("Notification status","Convert html successfully")
                    label_show = Label(root,text=response)
                elif mode == 'convertdata':
                    concert = editContent(str(current)).tailContent()
                    b = editContent(str(concert)).removeJunk()

                    c = editContent(str(b)).convertData()
                    c = c + take_CTA()

                    data_output_out.insert(END, c)

                    response = messagebox.showinfo("Notification status","Convert html successfully")
                    label_show = Label(root,text=response)
            def option_website():
                website = []
                for x in api.get_data("1X8sCZnp9ytXJBHz8-JC3tKMQZERCinzMiHCRKPtoiYY","Data Website!G5:G"):
                    try:
                        website.append(x[0])
                    except:
                        pass

                return website

            def take_CTA():
                CTA = []
                for x in api.get_data("1Z_Jy-8gX5lC0i_BoF4DnTpb7E2Ah9sypvV8eHDm-9LI","CTA!C3:C"):
                    try:
                        CTA.append(x[0])
                    except:
                        pass
                current = title_box.get("1.0", "end-1c")

                cta = random.choice(CTA)
                cta = cta.replace("{{domain}}","<code>" + display_selected()  + "</code>")
                cta = cta.replace("{{title}}","<code>" + current + "</code>")
                return cta

            def deleteLabel():
                data_output_out.delete("1.0",END)
                data_output.delete("1.0",END)

            def display_selected():
                choice = clicked.get()
                return choice

            def copy_to_clipboard():
                """Copy current contents of text_entry to clipboard."""
                root.clipboard_clear()  # Optional.
                root.clipboard_append(str(c))
                deleteLabel()
                response = messagebox.showinfo("Notification status", "Copy html successfully")
                label_show = Label(root, text=response)

            clicked = StringVar()
            clicked.set("Select website")
            drop = OptionMenu(root,clicked,*option_website(),command=display_selected)

            drop.place(x=450,y=15)

            title_box = Text(root,width=50,borderwidth=8,height = 2,bg = "light yellow")
            title_box.place(x= 50,y=10)

            data_output = Text(root,width=60,borderwidth=40,height = 20,bg = "light yellow")
            data_output.place(x= 50,y=70)

            data_output_out = Text(root,width=60,borderwidth=40,height = 20,bg = "light blue")
            data_output_out.place(x= 50,y=480)

            button1 = Button(root,text="Active",padx=15,pady=15,command=lambda: insertValue(mode='tailcontent'))
            button2 = Button(root,text="Convert",highlightbackground='white',padx=15,pady=15,command=lambda: insertValue(mode='convertdata'))
            button3 = Button(root,text="Copy",padx=15,pady=15,highlightbackground='red',command=lambda: copy_to_clipboard())
            button4 = Button(root,text="Delete",padx=15,pady=15,command=lambda: deleteLabel())

            button1.place(x=50,y=420)
            button2.place(x=130,y=420)
            button4.place(x=220,y=420)
            button3.place(x=485,y=420)

        if mode == 'choosefile':
            button_input = Button(root,text="Choose the path of the file",padx=40,pady=20,command= lambda: chooseFile()).grid(row=0,column=1)
            button_exit = Button(root,text="EXIT",padx=7,pady=7,command= root.quit).grid(row=0,column=2)
            data_output = Entry(root,width=45,borderwidth=15).grid(row=3,column=0,columnspan=3,padx=20,pady=20)

        elif mode == 'insertdata':
            insertData()

        root.mainloop()

class autoGUI():
    def __init__(self):
        pass
    def click(self,x,y):
        pyautogui.click()
        pass

if __name__ == "__main__":
    GUI().main_GUI(mode='insertdata')


