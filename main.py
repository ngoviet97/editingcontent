from bs4 import BeautifulSoup
import requests
import api
from tkinter import *
from tkinter import filedialog
from docx import Document
import docx
from tkinter import messagebox

root = Tk()
html1 = """<div>

Some of the objects<code>?M</code> button and formatting changes that you need to add to a document are easy to find in Google docs. But others can be quite difficult to locate, like a degree symbol. Our guide below will show you how to insert a degree symbol in Google Docs.
<h2>How to Get the Degrees Symbol on Google Docs</h2>
<div>
<div>

Use these steps when you need to put degrees symbols in your Google documents.

</div>
<code>Total Time:</code> 3 minutes

</div>
<code>Supply:</code>

<code>Tools:</code>

<code>Materials:</code> Google account login credentials

Sometimes you need to add special characters to a document that you are creating. It might be something like an arrow or bullet points, or superscript or subscript, but other situations could call for something a little less common, like a degree symbol. You might even want to know how to make a bullet point in Google Slides if you are working on a presentation.

Fortunately, Google Docs provides a way to insert a variety of special characters, and a degree symbol is one of the options that’s offered.

Our guide above has shown you how to add a degree symbol to Google Docs by following a few short steps.

The steps above were performed in the desktop version of the Google Chrome Web browser, but will work in other desktop browsers like Firefox or Edge, too.

Our tutorial continues below with additional discussion about how to insert a degree symbol in Google Docs.

Our Google Drive sign in tutorial can show you more about signing in to Google drive so that you can upload and access your files.
<h2>More Information on Using a Degree Symbol on Google Docs</h2>
There is a Windows keyboard shortcut for inserting a degree symbol as well. You can press <code>Alt + 0176</code> to add the symbol as well. Note that you need to use the numbers on your keyboard’s numeric keypad. It won’t work if you use the number row above the letters.

Symbols that you add to a document can be formatted in the same way as other text. You can highlight it with your mouse to select it, then you can do things like bold it, change the color of the text, change the size of the text, and more.

You could also choose to change the font style of the degree symbol, though that can cause it to look a little strange.

As you probably noticed when you were following our instructions on how to insert a degree symbol in Google Docs, there are a ton of other symbols available to you in Google Docs. So while finding the symbol and copying and pasting it from another location might be a good solution, you can also use the search feature for the symbols tool and see if it’s already available in Google’s word processing app.

Trying to make some formatting changes to emails that you are writing in Microsoft Outlook? Find out how to strikethrough in Outlook and easily draw a line through text.
<h2>Frequently Asked Questions About the Google Docs Degree Symbol</h2>
<div>
<div>
<div>
<h3>What’s the easiest way to insert degree symbol in a Google Docs document?</h3>
<div>

We have outlined a couple of different options for adding degree symbols above, but you might be looking for the easiest way to do this.

In my opinion, the easiest way to type degree symbols in a document is with the keyboard shortcut.

So if you need to include a degree symbol in a Google document, Google Sheets, or Google Slides, then using <code>Alt + 0176</code> on your keyboard’s numeric keypad is probably the fastest and simplest way to do this.

</div>
</div>
<div>
<h3>Is there a way to insert degree Celsius symbol?</h3>
<div>

While the steps in our article above have focused on ways for you to include a degree sign from the Insert menu in one of your documents in Google Docs, you might want to get a little more specific.

One additional change that you could want to use involves adding a Celsius modifier next to the degree symbol to indicate that you mean the temperature should be read as being in celsius.

You can find the Celsius degree symbol by going to <code>Insert &gt; Special characters &gt;</code> then clicking inside the search bar and typing the word <code>celsius</code>. You will then see a Celsius degree symbol that you can click to add to your document.

The same option works if you want to add a Fahrenheit degree symbol, too. You would just need to type the word <code>Fahrenheit</code> into the search bar. Or, since that’s a tough word to spell, just <code>fahr</code> should do the trick as well.

</div>
</div>
<div>
<h3>How can I add other mathematical symbols to my Google documents?</h3>
<div>

Our tutorial focuses specifically on the degree symbol, but almost all the symbols you might need to include in one of your documents can be found on the <code>Insert &gt; Special characters</code> drop down menu.

The exact option on that menu might vary based on the symbol that you need, so you might need to choose something other than the <code>Miscellaneous</code> that we selected to insert the degree symbol in our document.

</div>
</div>
<div>
<h3>What do I do if I can’t find a specific symbol in Google Docs?</h3>
<div>

If you’ve been looking through the different options on the Special characters menu and have even tried different options like Google Docs preferences of the Substitutions tab and you still can’t find the symbol you want, then the search bar in the Special characters menu can be really helpful.

If you go to <code>Insert &gt; Special characters</code> you will see a search bar on the right side of the dialog box. Simply type a description of the symbol you want and you should see a corresponding result if that character is available.

</div>
</div>
</div>
</div>
<h2>See also</h2>
<h3><em>Related</em></h3>
<div>
<table>
<tbody>
<tr>
<td style="width: 100%;">
<div style="width: 350px; padding-left: 5px; float: left; text-align: center;"><a href="http://eepurl.com/bwpdov" target="_blank" rel="nofollow noopener"><img src="http://update.downlater.net/wp-content/uploads/2022/06/1656097506_649_How-Do-I-Know-if-I-Have-Amazon-Prime.jpg" alt="solveyourtech.com newsletter" width="300" height="222" border="0" /><noscript><img src="http://update.downlater.net/wp-content/uploads/2022/06/1656097506_649_How-Do-I-Know-if-I-Have-Amazon-Prime.jpg" border="0" alt="solveyourtech.com newsletter" width="300" height="222"></noscript></a></div></td>
</tr>
</tbody>
</table>
<p style="text-align: center;">Disclaimer: Most of the pages on the internet include affiliate links, including some on this site.</p>

</div>
</div>"""

class editContent():
    def __init__(self,html):
        self.html = html
    def tailContent(self):
        tail_sheet = '1. Tail Content!C4:E'
        funct_sheet = '1. Tail Content!J19:J26'
        tag_sheet = '1. Tail Content!I29:I33'
        tail_id = '1Z_Jy-8gX5lC0i_BoF4DnTpb7E2Ah9sypvV8eHDm-9LI'

        dataset = []
        funct = []
        tag = []
        def createTuple(x,y,z):
            yield (x,y,z)
        def createSgTuple(x):
            yield (x)
        for x in api.get_data(tail_id,tail_sheet):
            try:
                tuple_create = createTuple(x[0],x[1],x[2])
                for x in tuple_create:
                    dataset.append(x)
            except:
                pass
        for y in api.get_data(tail_id,funct_sheet):
            try:
                tuple_funct = createSgTuple(y[0])
            except:
                pass
            for x in tuple_funct:
                funct.append(x)
        for z in api.get_data(tail_id,tag_sheet):
            try:
                tag_Craete = createSgTuple(z[0])
            except:
                pass
            for x in tag_Craete:
                tag.append(x)
        tag = tuple(tag)
        funct = tuple(funct)
        dataset = tuple(dataset)

        def replaceTail(data):
            for x in tag:
                for y in funct:
                    self.html = str(self.html).replace("</"+x+">" + str(data).strip(),"</"+x+">")
                    self.html = str(self.html).replace("</"+x+">" + y + str(data).strip(),"</"+x+">" + y)

            return self.html
        def setPriority(data):
            pri = [1,2,3,4]
            for k in pri:
                for z in data:
                    if z[2] == "Ưu tiên " + str(k) and z[1] == "Xoá bỏ":
                        replaceTail(z[0])
                    else:
                        pass

            return self.html

        return setPriority(dataset)
    def removeTag(self):
        spreadsheet_id = '1Z_Jy-8gX5lC0i_BoF4DnTpb7E2Ah9sypvV8eHDm-9LI'
        sheet_tags = '2. Tag!C4:G'

        def cr_tuple(x,y,z,t,h):
            yield (x,y,z,t,h)
        dataset = []
        for x in api.get_data(spreadsheet_id,sheet_tags):
            try:
                tuple = cr_tuple(x[0],x[1],x[2],x[3],x[4])
                for x in tuple:
                    dataset.append(x)
            except:
                pass

        def setPriority():
            def change_tag(html):
                data = []
                for l in html.splitlines():
                    if "pre>" not in l.strip():
                        value = ''.join('<p> {} </p>'.format(l.strip()))
                        data.append(value)
                    else:
                        value = l.strip()
                        data.append(value)
                xs = "\n".join(data)
                return xs
            change_html = change_tag(self.html)
            print(change_html)
            soup = BeautifulSoup(change_html, 'html.parser')
            def pri1(data):
                if data[4] == "Ưu tiên 1":
                    for x in soup.select(data[0]):
                        x.extract()
                else:
                    pass
            def pri2(data):
                if data[4] == "Ưu tiên 2":
                    pass
                else:
                    pass
            def pri3(data):
                if data[4] == "Ưu tiên 3":
                    for x in soup.findAll(data[0]):
                        try:
                            if data[1] != '' and data[2] in x[data[1]]:
                                x.extract()
                            elif data[1] == '' and data[2] in x.text:
                                x.extract()
                            else:
                                pass
                        except:
                            pass
            def pri4(data):
                if data[4] == "Ưu tiên 4" and data[0] == 'a':
                    for x in soup.findAll(data[0], href=True):
                        if x.text:
                            if data[1] == '' and data[2] in x.text:
                                x.unwrap()
                            elif data[1] != '' and data[2] in x[data[1]]:
                                x.unwrap()
                            elif len(x.text) > 80:
                                x.unwrap()
                            else:
                                pass
                    for x in soup.findAll(data[0], href=False):
                        if x.text:
                            x.unwrap()
                if data[4] == "Ưu tiên 4" and data[0] != 'a':
                    for x in soup.findAll(data[0]):
                        if x.text:
                            if data[1] == '' and data[2] in x.text:
                                x.unwrap()
                            elif data[1] != '' and data[2] in x[data[1]]:
                                x.unwrap()
                            elif len(x.text) > 80:
                                x.unwrap()
                            else:
                                pass
                return soup

            def remove_junk():
                delete_list = ['div', 'code', 'li','h2','h3','table','strong','p']
                [x.decompose() for x in soup.findAll(lambda tag: (not tag.contents or len(tag.get_text(strip=True)) <= 0) and tag.name == y for y in delete_list)]

                long_tag = ['strong','em','code','b']
                for y in long_tag:
                    for x in soup.findAll(y):
                        if len(x.text) > 27:
                            x.unwrap()
                        else:
                            pass
            def priority(data):
                pri1(data)
                pri2(data)
                pri3(data)
                pri4(data)
                remove_junk()

            for x in dataset:
                priority(x)

            soup = str(soup).replace("<code","<strong")

            soup = str(soup).replace("</code","</strong")

            return soup

        return setPriority()
    def removeJunk(self,data):
        spreadsheet_id = '1Z_Jy-8gX5lC0i_BoF4DnTpb7E2Ah9sypvV8eHDm-9LI'
        sheet_tags = '3. Delete!C4:D'

        def cr_tuple(x,y):
            yield (x,y)
        dataset = []
        for x in api.get_data(spreadsheet_id,sheet_tags):
            try:
                tuple = cr_tuple(x[0],x[1])
                for x in tuple:
                    dataset.append(x)
            except:
                pass

        for y in range(10):
            for x in dataset:
                data = str(data).replace(x[0],"")

        data = data.replace("<code></code>","")
        return data

class GUI():
    def __init__(self):
        root.title("EDIT CONTENT")
    def main_execution(self):
        output = editContent().main()
    def main_GUI(self):
        def exe_file():
            root.filename = filedialog.askopenfilename(initialdir="/Volumes/Root/Webify",title="Choose the path",filetypes=(("Word file","*.doc"),("Word file","*.docx"),("Note file","*.txt"),("all files","*.*")))
            doc = docx.Document(root.filename)
            all_paras = doc.paragraphs
            content = []
            for x in all_paras:
                content.append(x.text)
            content_trans = "\n".join(content)
            a = editContent(content_trans).tailContent()
            b = editContent(a).removeTag()
            print(b)
            list_b = b.split("\n")
            document = Document()
            for x in list_b:
                document.add_paragraph(x)
            document.save(root.filename)

            messagebox.showinfo("Status message","Successfully transfer")

        button_input = Button(root,text="Choose the path of the file",padx=40,pady=20,command= lambda: exe_file()).grid(row=0,column=1)
        button_exit = Button(root,text="EXIT",padx=7,pady=7,command= root.quit).grid(row=0,column=2)
        data_output = Entry(root,width=45,borderwidth=15).grid(row=3,column=0,columnspan=3,padx=20,pady=20)

        root.mainloop()
if __name__ == "__main__":
    GUI().main_GUI()


